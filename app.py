import uuid
import re
from pathlib import Path
from typing import List, Tuple, Optional

from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse

import pdfplumber
import fitz  # pymupdf
from docx import Document
from docx.shared import Inches

TMP_DIR = Path("/tmp")

STEP_ANCHORS = [
    ("Step 1 - Kapton position (View A)", "Fix the kapton in this position before assembly"),
    ("Step 2 - Before assembly: fill gel XTS-8030", "Before assembly: Fill volume"),
    ("Step 3 - Filling opening", "Openning for filling"),
    ("Step 4 - After assembly: fill gel XTS-8030", "After that MINID will be assembled"),
    ("Step 5 - Upper label area", "Area for label"),
]

app = FastAPI(title="WI from PDF API (Step Images + BOM)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # later you can restrict to your Netlify domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
def health():
    return {"status": "ok"}


@app.get("/")
def index():
    return {
        "service": "wi-from-pdf-api",
        "endpoints": ["/api/health", "/api/process-pdf", "/api/download/{filename}"],
    }


def extract_bom_rows_from_pdf(pdf_path: Path) -> List[Tuple[str, str, str, str]]:
    text_all: List[str] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for p in pdf.pages:
            text_all.append(p.extract_text() or "")
    text = "\n".join(text_all)

    bom: List[Tuple[str, str, str, str]] = []
    m = re.search(
        r"ITEM\s+QTY\.?\s+PART\s+NUMBER\s+DESCRIPTION(.*?)(?:Table\s+1|SIGNED\s+DATE|UNLESS\s+OTHERWISE|$)",
        text,
        flags=re.S | re.I,
    )
    if not m:
        return bom

    block = m.group(1)
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in block.splitlines()]
    for ln in lines:
        if not ln:
            continue
        mm = re.match(r"^(\d+)\s+(\d+)\s+([0-9A-Za-z.\-_]+)\s+(.*)$", ln)
        if mm:
            item, qty, part, desc = mm.groups()
            bom.append((item, qty, part, desc))

    return bom


def clip_from_anchor(
    page: fitz.Page,
    anchor_text: str,
    clip_above: int = 220,
    clip_below: int = 80,
    left_pad: int = 40,
    right_pad: int = 300,
) -> Optional[fitz.Rect]:
    rects = page.search_for(anchor_text)
    if not rects:
        return None

    r = rects[0]
    for rr in rects[1:]:
        r |= rr

    return fitz.Rect(
        max(0, r.x0 - left_pad),
        max(0, r.y0 - clip_above),
        min(page.rect.width, r.x1 + right_pad),
        min(page.rect.height, r.y1 + clip_below),
    )


def render_clip_to_png(
    pdf_path: Path,
    page_num: int,
    clip: fitz.Rect,
    out_path: Path,
    zoom: int = 2,
) -> Path:
    doc = fitz.open(str(pdf_path))
    page = doc.load_page(page_num)
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
    pix.save(str(out_path))
    doc.close()
    return out_path


def build_docx_with_step_images(pdf_path: Path, bom_rows, out_docx: Path) -> None:
    doc = Document()
    doc.add_heading("Work Instructions - Draft", 0)

    doc.add_heading("Safety", level=1)
    doc.add_paragraph("• PPE per procedures")
    doc.add_paragraph("• ESD protection")
    doc.add_paragraph("• Disconnect power before mechanical work")

    doc.add_heading("BOM (from drawing)", level=1)
    if bom_rows:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Qty"
        hdr[2].text = "Part Number"
        hdr[3].text = "Description"
        for item, qty, part, desc in bom_rows:
            row = table.add_row().cells
            row[0].text = str(item)
            row[1].text = str(qty)
            row[2].text = str(part)
            row[3].text = str(desc)
    else:
        doc.add_paragraph("BOM table not found.")

    doc.add_heading("Assembly process - step images", level=1)

    pdf = fitz.open(str(pdf_path))
    page = pdf.load_page(0)

    rendered_any = False
    for idx, (title, anchor) in enumerate(STEP_ANCHORS, 1):
        clip = clip_from_anchor(page, anchor)
        if clip is None:
            continue
        img_path = TMP_DIR / f"step_{idx}.png"
        render_clip_to_png(pdf_path, 0, clip, img_path, zoom=2)
        doc.add_heading(title, level=2)
        doc.add_picture(str(img_path), width=Inches(6.5))
        rendered_any = True

    pdf.close()

    if not rendered_any:
        doc.add_paragraph("No anchors found to crop images. PDF may be scanned (no selectable text).")

    doc.add_heading("Steps (summary)", level=1)
    for i, (title, _) in enumerate(STEP_ANCHORS, 1):
        doc.add_paragraph(f"{i}. {title}")

    doc.save(str(out_docx))


@app.post("/api/process-pdf")
async def process_pdf(
    request: Request,
    file: UploadFile = File(...),
    detail_level: str = Form("2"),
):
    if file.content_type not in ("application/pdf", "application/octet-stream"):
        raise HTTPException(status_code=400, detail="Please upload a PDF file")

    pdf_bytes = await file.read()
    pdf_path = TMP_DIR / f"pdf-{uuid.uuid4().hex}.pdf"
    pdf_path.write_bytes(pdf_bytes)

    bom_rows = extract_bom_rows_from_pdf(pdf_path)

    docx_name = f"wi-{uuid.uuid4().hex}.docx"
    docx_path = TMP_DIR / docx_name
    build_docx_with_step_images(pdf_path, bom_rows, docx_path)

    base_url = str(request.base_url).rstrip("/")
    docx_url = f"{base_url}/api/download/{docx_name}"

    return JSONResponse(
        {
            "received": {"filename": file.filename, "bytes": len(pdf_bytes), "detail_level": detail_level},
            "docx_url": docx_url,
        }
    )


@app.get("/api/download/{filename}")
def download_file(filename: str):
    path = TMP_DIR / filename
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        path=str(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
